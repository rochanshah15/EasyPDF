import os
from datetime import datetime
from flask import Flask, render_template, request, send_file, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
import PyPDF2
import pdfplumber
import fitz
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF
from PIL import Image
from pdf2docx import Converter
import pytesseract
from pdf2image import convert_from_path
from pdf2docx import Converter
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

# Get the absolute path for the directory where this app.py file lives
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# Check if running on Render and use its persistent disk path.
# Otherwise, create an 'instance' folder right next to your app.py file.
DATA_DIR = os.environ.get('RENDER_DATA_DIR', os.path.join(BASE_DIR, 'instance'))

# Create the data directory if it doesn't exist
os.makedirs(DATA_DIR, exist_ok=True)

app = Flask(__name__)

# Define paths for uploads and outputs within the data directory
UPLOAD_FOLDER = os.path.join(DATA_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(DATA_DIR, 'output')

# Configure app to use the new paths for file storage
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# Configure the absolute path for the database file
db_path = os.path.join(DATA_DIR, 'pdf_operations.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Ensure the subdirectories for uploads and outputs exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)


class PDFFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    uploaded_pdf = db.Column(db.String(255), nullable=False)
    manipulated_pdf = db.Column(db.String(255), nullable=True)
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

# Create database tables
with app.app_context():
    db.create_all()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/doc')
def doc():
    return render_template('doc.html')

@app.route('/merge', methods=['GET', 'POST'])
def merge():
    if request.method == 'POST':
        files = request.files.getlist('pdfs')
        if len(files) < 2:
            return "Please upload at least two PDF files."
        
        pdf_list = []
        uploaded_names = []
        for file in files:
            if file.filename == '':
                return "One of the files has no filename."
            if file and file.filename.lower().endswith('.pdf'):
                filename = secure_filename(file.filename)
                #upload folder
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                pdf_list.append(filepath)
                uploaded_names.append(filename)
        
        output_filename = 'merged.pdf'
        #output folder
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        merge_pdfs(pdf_list, output_path)

        # Store in database
        pdf_record = PDFFile(
            uploaded_pdf=','.join(uploaded_names),
            manipulated_pdf=output_filename
        )
        db.session.add(pdf_record)
        db.session.commit()
        
        return send_file(output_path, as_attachment=True)
    return render_template('merge.html')

import zipfile
import io
@app.route('/split', methods=['GET', 'POST'])
def split():
    if request.method == 'POST':
        file = request.files['pdf']
        if file.filename == '':
            return "No file selected."
        
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(filepath)

            # Split the PDF and get a list of in-memory split PDFs
            split_pages = split_pdf(filepath)

            # Create an in-memory ZIP file to store split PDFs
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for pdf_name, pdf_data in split_pages:
                    zip_file.writestr(pdf_name, pdf_data.read())  # Add each PDF to ZIP

            # Save to database (if needed)
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=f"split_{filename}.zip"
            )
            db.session.add(pdf_record)
            db.session.commit()

            # Return ZIP file as a response
            zip_buffer.seek(0)
            return send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name="split_pdfs.zip")

    return render_template('split.html')


@app.route('/extract-text', methods=['GET', 'POST'])
def extract_text_route():
    if request.method == 'POST':
        file = request.files['pdf']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = filename + '.txt'
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            extract_text(filepath, output_path)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('extract_text.html')


@app.route('/extract-images', methods=['GET', 'POST'])
def extract_images_route():
    if request.method == 'POST':
            file = request.files['pdf']
            if file.filename == '':
                return "No file selected."
            if file and file.filename.lower().endswith('.pdf'):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)

                # Extract images
                extracted_images = extract_images(filepath)

                if not extracted_images:
                    return "No images found in the PDF."

                # Convert extracted images to a PDF
                output_pdf_name = f"extracted_images_{filename}.pdf"
                output_pdf_path = os.path.join(app.config['OUTPUT_FOLDER'], output_pdf_name)
                images_to_pdf(extracted_images, output_pdf_path)

                return send_file(output_pdf_path, as_attachment=True)

    return render_template('extract_images.html')

@app.route('/extract-links', methods=['GET', 'POST'])
def extract_links_route():
    if request.method == 'POST':
        file = request.files['pdf']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = filename + '_links.txt'
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            extract_links(filepath, output_path)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('extract_links.html')

@app.route('/encrypt', methods=['GET', 'POST'])
def encrypt():
    if request.method == 'POST':
        file = request.files['pdf']
        password = request.form['password']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = 'encrypted_' + filename
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            encrypt_pdf(filepath, output_path, password)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('encrypt.html')

@app.route('/decrypt', methods=['GET', 'POST'])
def decrypt():
    if request.method == 'POST':
        file = request.files['pdf']
        password = request.form['password']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = 'decrypted_' + filename
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            decrypt_pdf(filepath, output_path, password)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('decrypt.html')

@app.route('/rearrange', methods=['GET', 'POST'])
def rearrange():
    if request.method == 'POST':
        file = request.files['pdf']
        page_order = request.form['page_order']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            page_order = [int(p) - 1 for p in page_order.split(',')]
            output_filename = 'rearranged_' + filename
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            rearrange_pages(filepath, output_path, page_order)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('rearrange.html')

@app.route('/rotate', methods=['GET', 'POST'])
def rotate():
    if request.method == 'POST':
        file = request.files['pdf']
        rotation = int(request.form['rotation'])
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = 'rotated_' + filename
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            rotate_pages(filepath, output_path, rotation)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('rotate.html')

@app.route('/add-metadata', methods=['GET', 'POST'])
def add_metadata_route():
    if request.method == 'POST':
        file = request.files['pdf']
        title = request.form['title']
        author = request.form['author']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = 'metadata_' + filename
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            add_metadata(filepath, output_path, title, author)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('add_metadata.html')

@app.route('/read-metadata', methods=['GET', 'POST'])
def read_metadata_route():
    if request.method == 'POST':
        file = request.files['pdf']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            metadata = read_metadata(filepath)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=None  # No manipulation, just reading
            )
            db.session.add(pdf_record)
            db.session.commit()

            return render_template('metadata_result.html', metadata=metadata)
    return render_template('read_metadata.html')

@app.route('/create-pdf', methods=['GET', 'POST'])
def create_pdf_route():
    if request.method == 'POST':
        text = request.form.get('text', 'No text provided')
        image = request.files.get('image')

        output_filename = 'created.pdf'  # Define the filename
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        image_path = None

        # Save the image if uploaded
        if image and image.filename:
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], image.filename)
            image.save(image_path)

        # Create the PDF with text and optional image
        create_pdf(output_path, text, image_path)
        
        # Store in database
        pdf_record = PDFFile(
            uploaded_pdf=image.filename if image and image.filename else 'text_only',
            manipulated_pdf=output_filename
        )
        db.session.add(pdf_record)
        db.session.commit()
        
        return send_file(output_path, as_attachment=True)

    return render_template('create_pdf.html')


@app.route('/optimize', methods=['GET', 'POST'])
def optimize():
    if request.method == 'POST':
        file = request.files['pdf']
        if file.filename == '':
            return "No file selected."
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_filename = 'optimized_' + filename
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            optimize_pdf(filepath, output_path)

            # Store in database
            pdf_record = PDFFile(
                uploaded_pdf=filename,
                manipulated_pdf=output_filename
            )
            db.session.add(pdf_record)
            db.session.commit()

            return send_file(output_path, as_attachment=True)
    return render_template('optimize.html')

@app.route('/convert', methods=['GET', 'POST'])
def convert_pdf_to_word():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            return "No file uploaded", 400  # Return HTTP 400 for bad requests

        file = request.files['pdf_file']
        if file.filename == '':
            return "No selected file", 400
        
        if file and file.filename.lower().endswith('.pdf'):
            pdf_filename = secure_filename(file.filename)  # Secure filename to prevent attacks
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
            file.save(pdf_path)

            # Generate output filename
            docx_filename = pdf_filename.rsplit('.', 1)[0] + ".docx"
            docx_path = os.path.join(app.config['OUTPUT_FOLDER'], docx_filename)

            try:
                # Convert PDF to Word
                cv = Converter(pdf_path)
                cv.convert(docx_path, start=0, end=None)
                cv.close()

                return send_file(docx_path, as_attachment=True)

            except Exception as e:
                return f"Error during conversion: {str(e)}", 500

        return "Invalid file format. Please upload a PDF.", 400
    
    return render_template('convert_pdf_to_word.html')

# PDF manipulation functions
def merge_pdfs(pdf_list, output_path):
    pdf_writer = PyPDF2.PdfWriter()
    for pdf in pdf_list:
        pdf_reader = PyPDF2.PdfReader(pdf)
        for page_num in range(len(pdf_reader.pages)):
            pdf_writer.add_page(pdf_reader.pages[page_num])
    with open(output_path, 'wb') as out:
        pdf_writer.write(out)

def split_pdf(pdf_path):
    """Splits a PDF into separate pages and returns a list of in-memory PDFs."""
    pdf_reader = PdfReader(pdf_path)
    split_files = []

    for i, page in enumerate(pdf_reader.pages):
        pdf_writer = PdfWriter()
        pdf_writer.add_page(page)
        
        # Create an in-memory buffer for the split page
        pdf_buffer = io.BytesIO()
        pdf_writer.write(pdf_buffer)
        pdf_buffer.seek(0)
        
        split_files.append((f"split_page_{i+1}.pdf", pdf_buffer))
    return split_files

def extract_text(pdf_path, output_txt_path):
    with pdfplumber.open(pdf_path) as pdf:
        full_text = ''
        for page in pdf.pages:
            full_text += page.extract_text() + '\n'
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write(full_text)

def extract_images(pdf_path):
    """Extracts images from a PDF file and returns a list of image file paths."""
    doc = fitz.open(pdf_path)
    image_paths = []

    for page_num in range(len(doc)):
        for img_index, img in enumerate(doc[page_num].get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]

            img_filename = f"image_{page_num+1}_{img_index+1}.{image_ext}"
            img_path = os.path.join(app.config['OUTPUT_FOLDER'], img_filename)

            with open(img_path, "wb") as img_file:
                img_file.write(image_bytes)

            image_paths.append(img_path)

    return image_paths


def images_to_pdf(image_paths, output_pdf_path):
    """Converts a list of image file paths into a single PDF without cropping images."""
    if not image_paths:
        return None  # No images extracted

    pdf = FPDF(unit="mm")  # Use mm for accurate scaling

    for image_path in image_paths:
        img = Image.open(image_path)
        img_width, img_height = img.size

        # Convert pixels to mm (assuming 1 inch = 25.4 mm and 1 inch = 96 pixels)
        mm_width = img_width * 25.4 / 96  # Convert pixels to mm
        mm_height = img_height * 25.4 / 96

        # Add new page with the same dimensions as the image
        pdf.add_page(format=(mm_width, mm_height))

        # Place image without resizing (so no cropping occurs)
        pdf.image(image_path, x=0, y=0, w=mm_width, h=mm_height)

    pdf.output(output_pdf_path, "F")
    return output_pdf_path


def extract_links(pdf_path, output_file):
    pdf_document = fitz.open(pdf_path)
    links = []
    for page_index in range(len(pdf_document)):
        page = pdf_document.load_page(page_index)
        link_list = page.get_links()
        for link in link_list:
            if link.get("uri"):
                links.append(link["uri"])
    
    with open(output_file, "w") as file:
        for link in links:
            file.write(link + "\n")
    return links

def encrypt_pdf(input_pdf, output_pdf, password):
    pdf_reader = PyPDF2.PdfReader(input_pdf)
    pdf_writer = PyPDF2.PdfWriter()
    for page_num in range(len(pdf_reader.pages)):
        pdf_writer.add_page(pdf_reader.pages[page_num])
    pdf_writer.encrypt(password)
    with open(output_pdf, 'wb') as out:
        pdf_writer.write(out)

def decrypt_pdf(input_pdf, output_pdf, password):
    pdf_reader = PyPDF2.PdfReader(input_pdf)
    pdf_reader.decrypt(password)
    pdf_writer = PyPDF2.PdfWriter()
    for page_num in range(len(pdf_reader.pages)):
        pdf_writer.add_page(pdf_reader.pages[page_num])
    with open(output_pdf, 'wb') as out:
        pdf_writer.write(out)

def rearrange_pages(input_pdf, output_pdf, page_order):
    pdf_reader = PyPDF2.PdfReader(input_pdf)
    pdf_writer = PyPDF2.PdfWriter()
    for page_num in page_order:
        pdf_writer.add_page(pdf_reader.pages[page_num])
    with open(output_pdf, 'wb') as out:
        pdf_writer.write(out)

def rotate_pages(input_pdf, output_pdf, rotation):
    pdf_reader = PyPDF2.PdfReader(input_pdf)
    pdf_writer = PyPDF2.PdfWriter()
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        page.rotate(rotation)
        pdf_writer.add_page(page)
    with open(output_pdf, 'wb') as out:
        pdf_writer.write(out)

def add_metadata(input_pdf, output_pdf, title, author):
    pdf_reader = PyPDF2.PdfReader(input_pdf)
    pdf_writer = PyPDF2.PdfWriter()
    for page_num in range(len(pdf_reader.pages)):
        pdf_writer.add_page(pdf_reader.pages[page_num])
    metadata = {
        '/Title': title,
        '/Author': author,
        '/Producer': ''
    }
    pdf_writer.add_metadata(metadata)
    with open(output_pdf, 'wb') as out:
        pdf_writer.write(out)

def read_metadata(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    return pdf_reader.metadata

# Create a PDF directly from the text.
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
def create_pdf(output_pdf, text, image_path=None):
    """
    Creates a PDF with the provided text and an optional image.
    
    :param output_pdf: Path to save the generated PDF.
    :param text: Text to include in the PDF.
    :param image_path: Path to an image file (optional).
    """
    c = canvas.Canvas(output_pdf, pagesize=letter)
    
    # Add user text
    c.drawString(100, 750, text)
    
    # Add image if provided
    if image_path:
        try:
            c.drawImage(image_path, 100, 500, width=200, height=150)
        except Exception as e:
            print(f"Error adding image: {e}")

    c.save()

def optimize_pdf(input_file, output_file):
    pdf_document = fitz.open(input_file)
    pdf_document.save(output_file, garbage=3, deflate=True)

# def convert_pdf_to_word(input_pdf, output_docx):
#     cv = Converter(input_pdf)  # Create a converter object
#     cv.convert(output_docx, start=0, end=None)  # Convert entire PDF
#     cv.close()  # Close the converter

def is_scanned_pdf(pdf_path):
    """Check if PDF is scanned (no selectable text)."""
    doc = fitz.open(pdf_path)
    for page in doc:
        text = page.get_text("text")
        if text.strip():  # If text exists, it's not scanned
            return False
    return True

def convert_pdf_to_word(pdf_path, word_path):
    """Convert any PDF (normal or scanned) to Word."""
    if is_scanned_pdf(pdf_path):
        print("Detected a scanned PDF. Using OCR...")
        convert_scanned_pdf_to_word(pdf_path, word_path)
    else:
        convert_normal_pdf_to_word(pdf_path, word_path)

def convert_normal_pdf_to_word(pdf_path, word_path):
    """Convert a regular (text-based) PDF to Word using pdf2docx."""
    try:
        cv = Converter(pdf_path)
        cv.convert(word_path, start=0, end=None)
        cv.close()
        print("PDF successfully converted to Word!")
    except Exception as e:
        print(f"pdf2docx failed: {e}")
        extract_text_with_pymupdf(pdf_path, word_path)

def extract_text_with_pymupdf(pdf_path, word_path):
    """Extract text from a normal PDF using PyMuPDF."""
    doc = fitz.open(pdf_path)
    word_doc = Document()

    for page in doc:
        text = page.get_text("text")
        word_doc.add_paragraph(text)

    word_doc.save(word_path)

# def convert_scanned_pdf_to_word(pdf_path, word_path):
#     """Convert a scanned (image-based) PDF to Word using OCR."""
#     images = convert_from_path(pdf_path)
#     doc = Document()

#     for img in images:
#         text = pytesseract.image_to_string(img)
#         doc.add_paragraph(text)

#     doc.save(word_path)

import pytesseract
from pdf2image import convert_from_path
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter

def preprocess_image(image):
    """Enhance image quality for better OCR accuracy."""
    image = image.convert('L')  # Convert to grayscale
    image = image.filter(ImageFilter.SHARPEN)  # Sharpen the image
    image = ImageEnhance.Contrast(image).enhance(2)  # Increase contrast
    return image

def convert_scanned_pdf_to_word(pdf_path, word_path):
    """Convert a scanned (image-based) PDF to Word using OCR."""
    images = convert_from_path(pdf_path)
    doc = Document()
    
    for i, img in enumerate(images):
        img = preprocess_image(img)  # Preprocess for better OCR
        text = pytesseract.image_to_string(img).strip()

        if text:  # Only add non-empty pages
            doc.add_paragraph(text)
            doc.add_page_break()  # Add a page break for formatting

    if len(doc.paragraphs) == 0:
        doc.add_paragraph("No text detected in the document.")

    doc.save(word_path)
    print(f"Conversion completed: {word_path}")



if __name__ == "__main__":
    with app.app_context():  # ✅ FIX: Set application context
        app.run(debug=True, port=5500)